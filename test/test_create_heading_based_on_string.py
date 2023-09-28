import pytest
from utils import create_headings

from docx import Document

#Creating a temporary document.
@pytest.fixture
def temp_document():
    doc = Document()
    yield doc

# Define test cases
@pytest.mark.parametrize("input_heading, expected_heading, expected_level", [
    ('1.1.1 Test description','1.1.1 Test description', 3),
    ('1.1 Test description','1.1 Test description', 2),
    ('11 Test description','11 Test description', 1),
    ])

def test_input_heading(temp_document, input_heading, expected_heading, expected_level)->None:
    """
    Ensure that headings can be added as expected without any truncations or unexpected characters.
    """
    (text, level) = create_headings.add_heading(document=temp_document, text=input_heading)
    # Extract the text from the document
    doc_text = '\n'.join(paragraph.text for paragraph in temp_document.paragraphs)
    assert doc_text == expected_heading
    assert (text, level) == (expected_heading, expected_level)


