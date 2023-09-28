from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from utils import list_files



def add_heading(document: Document, text: str):
    """
    format string based on name of directory
    """
    level = len(text.split(' ')[0].split('.'))
    heading = document.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = heading.add_run(text)
    run.bold = True
    run.style.font.size = 14 * 2
    run.style.font.name = 'Arial'
    run.style.font.color.rgb = RGBColor(0, 0, 0)
    heading.style = f'Heading {level}'
    return (text, level)


def create_headings_for_file(filePath: str, rootPath: str, createdHeadings: list[str], document: Document)->list[str]:
    """
    Format the string according to the directory name and its position, along with the subdirectory index.
    """
    if rootPath[-1:] != '/':
        rootPath = f'{rootPath}/'
    filePath = filePath.replace(rootPath, './', 1)
    splittedFilePath = filePath.split('/')
    for deepness, heading in enumerate(splittedFilePath):
        if len(heading)>0:
            if heading[0].isdigit():
                splitted_heading = heading.split('_')
                subindex=''
                for prev_index in range(deepness):
                    if splittedFilePath[prev_index][0].isdigit():
                        subindex=(f'{subindex}.{splittedFilePath[prev_index].split("_")[0]}')
                name= (f'{subindex}.{" ".join(item for item in splitted_heading)}'[1:])
                if (name not in createdHeadings) and ('.' not in heading):
                    add_heading(document=document, text=name)
                    createdHeadings.append(name)
    return createdHeadings   
